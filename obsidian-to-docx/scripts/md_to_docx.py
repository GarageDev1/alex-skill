#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
obsidian-to-docx · 核心转换器
把一个 Obsidian markdown(含 ECharts 代码块 / wikilink / 表格)导出成好看的 .docx。

设计原则
- 这是一个 exporter,不改写内容,只做"markdown 元素 → docx 元素"的映射。
- ECharts 代码块在这里只负责"定位 + 占位/嵌图";真正的 JSON→PNG 渲染交给
  render_echarts.js(puppeteer)。本脚本通过 --charts-dir 接收已渲染好的 PNG。
- 样式由 profile 决定(default 通用 / equity-research 研报封面),见 references/style-profiles.md。

用法
  python md_to_docx.py --input <a.md> --output <a.docx>
  python md_to_docx.py --input <a.md> --output <a.docx> --charts-dir <png目录> --profile equity-research

版式要点(M2.5 样式层)
- 标题:微软雅黑 + 主题色,H1 带底部分隔线,段前段后间距 + keep_with_next
- 正文:宋体/Times 双字体 10.5pt,1.4 倍行距,段后 6pt,两端对齐
- 表格:三线表(顶/底主题色粗线 + 表头底纹,无竖线,数据行统一左对齐)
- 引用:普通引用 → 左侧主题色条 + 浅灰底纹 callout;📊/📂 图注 → 小字灰色 caption
- 代码块:浅灰底纹 + 细边框,保留换行
- 页眉(文档标题)/ 页脚(页码域)
"""

import argparse
import os
import re
import sys

from urllib.parse import unquote

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- 字体 / 样式常量 ----------
ZH_FONT = "宋体"               # 正文中文(east-asia)
ZH_HEADING_FONT = "微软雅黑"    # 标题中文(标题黑体、正文宋体是中文文档惯例)
EN_FONT = "Times New Roman"    # 正文英文/数字(ascii + hAnsi)
BODY_SIZE = Pt(10.5)           # 五号
PLACEHOLDER_GRAY = RGBColor(0x99, 0x99, 0x99)
CAPTION_GRAY = RGBColor(0x59, 0x59, 0x59)
DARK_TEXT = RGBColor(0x26, 0x26, 0x26)

# 主题色按 profile 切换:default 深蓝 / equity-research 研报红
THEMES = {
    "default": {"accent": "1F4E79", "header_fill": "DEEBF7"},
    "equity-research": {"accent": "C00000", "header_fill": "F2DCDB"},
}
THEME = THEMES["default"]  # main() 按 --profile 重设


def _rgb(hexstr):
    return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))


# ============================================================
# 双字体 helper:让一个 run 中英文各用各的字体
# ============================================================
def set_run_fonts(run, zh=ZH_FONT, en=EN_FONT, size=None):
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), en)
    rfonts.set(qn("w:hAnsi"), en)
    rfonts.set(qn("w:eastAsia"), zh)
    if size is not None:
        run.font.size = size


# ============================================================
# 底纹 / 边框 低层 helper
# ============================================================
def _set_shading(pr_element, fill):
    """给 pPr / tcPr 加底纹"""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pr_element.append(shd)


def _set_para_borders(p, edges):
    """段落边框。edges = {"left": (sz, color), "bottom": (...)};sz 单位 1/8 pt"""
    ppr = p._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    for edge, (sz, color) in edges.items():
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "4")
        e.set(qn("w:color"), color)
        pbdr.append(e)


def _tight(paragraph, line=1.0, after=0):
    """收紧段落(表格单元格/页眉脚等不吃正文的行距和段后距)"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line
    pf.space_after = Pt(after)
    pf.space_before = Pt(0)


# ============================================================
# 超链接 helper(图注来源、附录引用)
# ============================================================
def add_hyperlink(paragraph, url, text, size=None):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), EN_FONT)
    rfonts.set(qn("w:hAnsi"), EN_FONT)
    rfonts.set(qn("w:eastAsia"), ZH_FONT)
    rpr.append(rfonts)
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size.pt * 2)))
        rpr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


# ============================================================
# LaTeX → OMML(Word 原生公式)
# latex2mathml 转 MathML,再用 Office 自带 MML2OMML.XSL 转 OMML。
# 任一环节失败返回 None,调用方退回纯文本,导出不中断。
# ============================================================
_MML2OMML_CANDIDATES = [
    r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
    r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL",
]
_omml_xslt = None


def _get_omml_xslt():
    global _omml_xslt
    if _omml_xslt is None:
        from lxml import etree
        for cand in _MML2OMML_CANDIDATES:
            if os.path.exists(cand):
                _omml_xslt = etree.XSLT(etree.parse(cand))
                break
        if _omml_xslt is None:
            _omml_xslt = False  # 标记不可用
    return _omml_xslt


def latex_to_omml(latex, display=False):
    """LaTeX 字符串 → OMML lxml 元素(m:oMathPara 或 m:oMath);失败返回 None"""
    try:
        import latex2mathml.converter
        from lxml import etree
        xslt = _get_omml_xslt()
        if not xslt:
            return None
        mathml = latex2mathml.converter.convert(
            latex, display="block" if display else "inline")
        omml = xslt(etree.fromstring(mathml)).getroot()
        if not display:
            # 行内取 oMath 本体(剥掉 oMathPara 包装,避免独占一行)
            M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
            if omml.tag == M + "oMathPara":
                inner = omml.find(M + "oMath")
                if inner is not None:
                    omml = inner
        return omml
    except Exception:
        return None


# 行内 $...$ 是否真公式(防货币误判:HK$1,073 / $88.34 一律不转):
# 必须含 LaTeX 特征字符——反斜杠命令,或 ^/_ 上下标(配合 {})
_MATH_HINT = re.compile(r"\\[A-Za-z]+|[\^_]\{|[\^_][A-Za-z0-9]")


def _looks_like_math(content):
    return bool(_MATH_HINT.search(content))


# ============================================================
# inline 解析:**粗** *斜* `码` [文字](url) [[wikilink]] $公式$
# 返回 [(text, {bold,italic,code,url,math}), ...]
# ============================================================
INLINE_RE = re.compile(
    r"(\*\*.+?\*\*)"            # bold
    r"|(\*.+?\*)"              # italic
    r"|(`[^`]+`)"             # inline code
    r"|(\[\[.+?\]\])"         # wikilink
    r"|(\[.+?\]\(.+?\))"      # md link
    r"|(\$[^$\n*`]+\$)"       # inline math(候选,需过 _looks_like_math;内部排除 * 和 ` 防止货币 $..$ 跨吞加粗/斜体/代码标记)
)

# [[页面|显示名]] → 显示名 / [[页面]] → 页面;粗体/斜体内部嵌套的双链同样要清掉
_WIKILINK_RE = re.compile(r"\[\[([^\]]+?)\]\]")


def _strip_wikilinks(s):
    return _WIKILINK_RE.sub(
        lambda m: m.group(1).split("|", 1)[1] if "|" in m.group(1) else m.group(1), s
    )


# ============================================================
# 图片 embed:![[name]] (Obsidian) / ![](url|path) (markdown)
# 整行(可含前导列表符)由一个/多个图片 embed 构成 → 当作图片块
# ============================================================
_IMG_EMBED = r"(?:!\[\[[^\]]+\]\]|!\[[^\]]*\]\([^)]+\))"
IMAGE_LINE_RE = re.compile(r"^\s*(?:[-*+]\s+)?(?:" + _IMG_EMBED + r")+\s*$")
IMG_EXTRACT_RE = re.compile(_IMG_EMBED)


def parse_embed(token):
    """单个图片 embed token → {kind, target}"""
    token = token.strip()
    if token.startswith("![["):
        inner = token[3:-2]
        target = inner.split("|", 1)[0].strip()  # 去掉 |尺寸
        return {"kind": "wiki", "target": target}
    m = re.match(r"!\[[^\]]*\]\(([^)]+)\)", token)
    return {"kind": "md", "target": m.group(1).strip()}


def parse_inline(text):
    out = []
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            out.append((text[pos:m.start()], {}))
        tok = m.group(0)
        if tok.startswith("**"):
            out.append((_strip_wikilinks(tok[2:-2]), {"bold": True}))
        elif tok.startswith("*"):
            out.append((_strip_wikilinks(tok[1:-1]), {"italic": True}))
        elif tok.startswith("`"):
            out.append((tok[1:-1], {"code": True}))
        elif tok.startswith("[["):
            out.append((_strip_wikilinks(tok), {}))
        elif tok.startswith("$"):
            inner = tok[1:-1]
            if _looks_like_math(inner):
                out.append((inner, {"math": True}))
            else:  # 货币等非公式,原样保留含 $ 的文本
                out.append((tok, {}))
        else:  # [text](url)
            mm = re.match(r"\[(.+?)\]\((.+?)\)", tok)
            out.append((mm.group(1), {"url": mm.group(2)}))
        pos = m.end()
    if pos < len(text):
        out.append((text[pos:], {}))
    return out


def add_inline_runs(paragraph, text, size=BODY_SIZE, color=None):
    for seg, fmt in parse_inline(text):
        if not seg:
            continue
        if fmt.get("url"):
            add_hyperlink(paragraph, fmt["url"], seg, size=size)
            continue
        if fmt.get("math"):
            omml = latex_to_omml(seg, display=False)
            if omml is not None:
                paragraph._p.append(omml)
                continue
            # 转换失败 → 退回原文(带 $ 提示未渲染)
            seg = "$" + seg + "$"
        run = paragraph.add_run(seg)
        run.bold = fmt.get("bold", False)
        run.italic = fmt.get("italic", False)
        if fmt.get("code"):
            set_run_fonts(run, zh="Consolas", en="Consolas", size=size)
        else:
            set_run_fonts(run, size=size)
        if color is not None:
            run.font.color.rgb = color


# ============================================================
# 块级解析:把 md 文本拆成 block 列表
# ============================================================
def strip_frontmatter(text):
    if text.startswith("---"):
        m = re.match(r"^---\s*\n(.*?)\n---\s*\n", text, re.DOTALL)
        if m:
            return m.group(1), text[m.end():]
    return "", text


def parse_frontmatter(fm_text):
    """轻量 YAML:只解析顶层 key: value(够 rating box 用,不处理嵌套/列表)"""
    d = {}
    for line in fm_text.split("\n"):
        m = re.match(r"^([A-Za-z_][\w-]*):\s*(.*)$", line)
        if m:
            d[m.group(1)] = m.group(2).strip()
    return d


def parse_blocks(body):
    """返回 block 列表,每个是 dict {type, ...}"""
    blocks = []
    lines = body.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            i += 1
            continue

        # 水平分隔线 --- / *** / ___
        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", stripped):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # 代码块(含 echarts)
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            buf = []
            code_start = i + 1  # 1-based 起始行(用于未闭合警告)
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            if i >= n:
                # 代码块未闭合:从开栅栏到文件尾全被当代码吞掉,
                # 是"导出后大片正文/标题莫名缺失"的常见原因。打 stderr 警告,不静默。
                sys.stderr.write(
                    f"[警告] 第 {code_start} 行的代码块(```{lang or ''})未闭合,"
                    f"其后到文末的内容被按代码处理,可能丢失标题/正文。\n"
                )
            else:
                i += 1  # 跳过结束 ```
            if lang == "echarts":
                blocks.append({"type": "echarts", "code": "\n".join(buf)})
            else:
                blocks.append({"type": "code", "lang": lang, "code": "\n".join(buf)})
            continue

        # 标题(ATX):行首允许 ≤3 个空格/tab 缩进(CommonMark 合规;Obsidian 同样渲染为标题)。
        # 原来锚 ^ 不吃前导空白,导致 "  ## 标题" / "\t## 标题" 落进段落分支被吞掉,导出后标题丢失。
        # 4 个空格仍不识别(CommonMark 里 4 空格是缩进代码块,不是标题)。
        m = re.match(r"^[ \t]{0,3}(#{1,6})[ \t]+(.*)$", line)
        if m:
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2).strip()})
            i += 1
            continue

        # 表格(当前行含 | 且下一行是分隔行)
        if "|" in line and i + 1 < n and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]):
            tbl = [line]
            i += 1  # 表头
            tbl.append(lines[i])  # 分隔
            i += 1
            while i < n and "|" in lines[i] and lines[i].strip():
                tbl.append(lines[i])
                i += 1
            blocks.append({"type": "table", "rows": tbl})
            continue

        # 块级公式 $$...$$(单行或多行)
        if stripped.startswith("$$"):
            if stripped.endswith("$$") and len(stripped) > 4:  # 单行 $$...$$
                blocks.append({"type": "math", "latex": stripped[2:-2].strip()})
                i += 1
                continue
            buf = [stripped[2:]]
            i += 1
            while i < n and "$$" not in lines[i]:
                buf.append(lines[i])
                i += 1
            if i < n:
                buf.append(lines[i].split("$$")[0])
                i += 1
            blocks.append({"type": "math", "latex": "\n".join(buf).strip()})
            continue

        # 引用块(图注 > 也走这里)
        if stripped.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip()[1:].strip())
                i += 1
            blocks.append({"type": "quote", "lines": buf})
            continue

        # 图片行(整行是图片 embed,可含前导列表符;吞掉后续连续图片行并成一块)
        if IMAGE_LINE_RE.match(line):
            imgs = [parse_embed(m.group(0)) for m in IMG_EXTRACT_RE.finditer(line)]
            i += 1
            while i < n and IMAGE_LINE_RE.match(lines[i]):
                imgs += [parse_embed(m.group(0)) for m in IMG_EXTRACT_RE.finditer(lines[i])]
                i += 1
            blocks.append({"type": "images", "items": imgs})
            continue

        # 列表
        if re.match(r"^\s*([-*+]|\d+\.)\s+", line):
            items = []
            while i < n and re.match(r"^\s*([-*+]|\d+\.)\s+", lines[i]):
                lm = re.match(r"^\s*([-*+]|\d+\.)\s+(.*)$", lines[i])
                ordered = bool(re.match(r"^\s*\d+\.", lines[i]))
                items.append({"ordered": ordered, "text": lm.group(2)})
                i += 1
            blocks.append({"type": "list", "items": items})
            continue

        # Setext 标题:普通文本行紧跟 === (H1) 或 --- (H2) 下划线。
        # 这是 markdown 的第二种标题语法,Obsidian 渲染为标题,但原解析器没覆盖、当成段落,导出后标题丢失。
        # 边界:无空行隔开的 "段落\n---" 在 CommonMark 里本就是 Setext H2(与 Obsidian 一致);
        # 若本意是水平分隔线,应在 --- 前留空行。--- 要求至少 3 个,与 hr 阈值统一,避免 "正文\n-" 误判。
        if i + 1 < n:
            nxt = lines[i + 1].strip()
            if re.fullmatch(r"=+", nxt):
                blocks.append({"type": "heading", "level": 1, "text": stripped})
                i += 2
                continue
            if re.fullmatch(r"-{3,}", nxt):
                blocks.append({"type": "heading", "level": 2, "text": stripped})
                i += 2
                continue

        # 普通段落(连续非空行合并)
        buf = [stripped]
        i += 1
        while i < n and lines[i].strip() and not _is_block_start(lines[i], lines, i):
            buf.append(lines[i].strip())
            i += 1
        blocks.append({"type": "para", "text": " ".join(buf)})
    return blocks


def _is_block_start(line, lines, i):
    s = line.strip()
    if s.startswith(("#", ">", "```", "$$")):
        return True
    if IMAGE_LINE_RE.match(line):
        return True
    if re.match(r"^\s*([-*+]|\d+\.)\s+", line):
        return True
    if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]):
        return True
    return False


# ============================================================
# 表格行解析
# ============================================================
def split_row(row):
    # 按未转义的 | 切列(表格内 wikilink 显示名写作 [[页面\|显示名]]),切完再还原 \| → |
    row = row.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|") and not row.endswith("\\|"):
        row = row[:-1]
    cells = re.split(r"(?<!\\)\|", row)
    return [c.strip().replace("\\|", "|") for c in cells]


# 数字单元格(可带货币符/千分位/百分号/倍数)→ 右对齐
_NUM_CELL_RE = re.compile(r"^[\$¥€\+\-]?\s*[\d,]+(\.\d+)?\s*([%xX倍]|bp|pct)?$")


def _is_numeric_cell(text):
    plain = re.sub(r"(\*\*|\*|`)", "", text).strip()
    return bool(plain) and bool(_NUM_CELL_RE.match(plain))


# ============================================================
# 渲染:block → docx
# ============================================================
HEADING_SIZES = {1: Pt(18), 2: Pt(15), 3: Pt(13), 4: Pt(12), 5: Pt(11), 6: Pt(11)}
HEADING_SPACING = {1: (18, 10), 2: (14, 8), 3: (10, 6), 4: (8, 4)}  # (段前, 段后) pt


def render_heading(doc, level, text):
    lvl = min(level, 4)
    p = doc.add_heading(level=lvl)
    p.text = ""  # 清掉 add_heading 默认 run
    run = p.add_run(text)
    # 标题中英文统一微软雅黑(标题黑体、正文宋体的中文排版惯例)
    set_run_fonts(run, zh=ZH_HEADING_FONT, en=ZH_HEADING_FONT, size=HEADING_SIZES.get(level, Pt(12)))
    run.bold = True
    # H1/H2 主题色,H3/H4 近黑
    run.font.color.rgb = _rgb(THEME["accent"]) if lvl <= 2 else DARK_TEXT
    before, after = HEADING_SPACING[lvl]
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.keep_with_next = True  # 标题不孤悬页底
    if lvl == 1:
        _set_para_borders(p, {"bottom": (8, THEME["accent"])})  # 1pt 主题色底线
    return p


def render_blocks(doc, blocks, charts_dir=None, md_dir=".", vault_root="."):
    echarts_idx = 0
    img_cache = {"idx": None}
    for b in blocks:
        t = b["type"]
        if t == "heading":
            render_heading(doc, b["level"], b["text"])

        elif t == "math":
            omml = latex_to_omml(b["latex"], display=True)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if omml is not None:
                p._p.append(omml)
            else:
                run = p.add_run("$$" + b["latex"] + "$$")
                set_run_fonts(run, zh="Consolas", en="Consolas", size=Pt(9))

        elif t == "para":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline_runs(p, b["text"])

        elif t == "list":
            for it in b["items"]:
                style = "List Number" if it["ordered"] else "List Bullet"
                p = doc.add_paragraph(style=style)
                p.paragraph_format.space_after = Pt(3)
                add_inline_runs(p, it["text"])

        elif t == "quote":
            render_quote(doc, b["lines"])

        elif t == "table":
            render_table(doc, b["rows"])

        elif t == "code":
            render_code(doc, b["code"])

        elif t == "echarts":
            echarts_idx += 1
            render_echarts_block(doc, echarts_idx, b["code"], charts_dir)

        elif t == "images":
            render_images(doc, b["items"], md_dir, vault_root, img_cache)

        elif t == "hr":
            add_hr(doc)


def render_quote(doc, quote_lines):
    """两种引用:📊/📂 图注 → 小字灰 caption;普通引用 → 左色条 + 浅灰底 callout"""
    is_caption = all(ln.startswith(("📊", "📂")) for ln in quote_lines if ln)
    if is_caption:
        for ln in quote_lines:
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Inches(0.25)
            pf.space_after = Pt(2)
            pf.line_spacing = 1.15
            add_inline_runs(p, ln, size=Pt(9), color=CAPTION_GRAY)
        # 图注组后补一点间隔
        doc.paragraphs[-1].paragraph_format.space_after = Pt(8)
        return
    # callout:相邻段落边框/底纹一致时 Word 自动连成一个框
    for k, ln in enumerate(quote_lines):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.15)
        pf.line_spacing = 1.3
        pf.space_before = Pt(0)
        pf.space_after = Pt(8 if k == len(quote_lines) - 1 else 2)
        _set_para_borders(p, {"left": (24, THEME["accent"])})  # 3pt 主题色条
        _set_shading(p._p.get_or_add_pPr(), "F7F7F7")
        add_inline_runs(p, ln, size=Pt(9.5), color=RGBColor(0x40, 0x40, 0x40))


def render_code(doc, code):
    """代码块:浅灰底纹 + 细边框,保留换行(原实现 \\n 塞一个 run 会被 Word 折叠)"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(8)
    pf.left_indent = Inches(0.1)
    _set_shading(p._p.get_or_add_pPr(), "F5F5F5")
    _set_para_borders(p, {e: (4, "E0E0E0") for e in ("top", "bottom", "left", "right")})
    run = p.add_run()
    for k, ln in enumerate(code.split("\n")):
        if k:
            run.add_break()
        run.add_text(ln)
    set_run_fonts(run, zh="Consolas", en="Consolas", size=Pt(9))


def add_hr(doc):
    """水平分隔线 = 带底边框的空段落"""
    p = doc.add_paragraph()
    _tight(p)
    _set_para_borders(p, {"bottom": (6, "CCCCCC")})
    return p


def style_table_threeline(table, accent=None, header_fill=None):
    """三线表:顶/底 1.5pt 主题色,行间 0.5pt 浅灰,无竖线"""
    accent = accent or THEME["accent"]
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    spec = {
        "top": ("single", "12", accent),
        "bottom": ("single", "12", accent),
        "insideH": ("single", "4", "DDDDDD"),
        "left": ("none", "0", "auto"),
        "right": ("none", "0", "auto"),
        "insideV": ("none", "0", "auto"),
    }
    for edge, (val, sz, color) in spec.items():
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), val)
        if val != "none":
            e.set(qn("w:sz"), sz)
            e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


def _shade_cell(cell, fill):
    _set_shading(cell._tc.get_or_add_tcPr(), fill)


def _content_width_in(doc):
    """正文可用宽度(英寸):页宽 - 左右页边距;取不到回退 6.3"。"""
    try:
        s = doc.sections[0]
        w_in = (s.page_width - s.left_margin - s.right_margin) / 914400.0
        if 3.0 < w_in < 12.0:
            return w_in
    except Exception:
        pass
    return 6.3


def _cell_visual_len(s):
    """估算单元格显示宽度:剥 markdown inline 标记后,中文/全角=2、ascii=1。"""
    s = re.sub(r"(\*\*|\*|`|\[\[|\]\]|\[|\]\([^)]*\))", "", s)
    return sum(2 if ord(c) > 0x2E80 else 1 for c in s)


def _compute_col_widths(header, data, total_in, max_ratio=0.6):
    """按各列内容长度比例分配列宽(英寸)。

    python-docx 建表默认 gridCol 等分,长短不一的列被拉成等宽、长文本列折行严重。
    每列按"最长单元格的视觉宽度"定权重,clamp 到 [min_in, total*max_ratio] 后归一化回 total,
    让内容多的列宽、内容少的列窄。

    min_in 随列数自适应:列少时 0.7"(够放表头),列多时下沉到 0.35",
    否则 min 会把大多数列 clamp 成等宽、失去按内容区分的能力。
    """
    ncol = len(header)
    if ncol == 0:
        return []
    min_in = max(0.35, min(0.7, total_in / ncol * 0.55))
    weights = []
    for j in range(ncol):
        cells = [header[j]] + [r[j] if j < len(r) else "" for r in data]
        # 单元格会自动折行,最长格按 ~80 视觉宽封顶,避免一个超长格独占整张表
        wl = [min(_cell_visual_len(c), 80) for c in cells]
        weights.append(max(wl) if wl else 1)
    if sum(weights) == 0:
        weights = [1] * ncol
    raw = [total_in * w / sum(weights) for w in weights]
    clamped = [max(min_in, min(w, total_in * max_ratio)) for w in raw]
    s = sum(clamped)
    return [w * total_in / s for w in clamped]


def _set_col_widths(table, widths_in):
    """显式设定列宽:重写 gridCol + 每个 cell 的 tcW + fixed 布局。

    fixed 布局让 Word 严格尊重设定列宽,不被 autofit 改写回等宽。
    """
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for old in grid.findall(qn("w:gridCol")):
        grid.remove(old)
    for w in widths_in:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 1440)))
        grid.append(gc)
    tblPr = tbl.tblPr
    for tag in ("w:tblLayout", "w:tblW"):
        for old in tblPr.findall(qn(tag)):
            tblPr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(sum(widths_in) * 1440)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j < len(widths_in):
                tcPr = cell._tc.get_or_add_tcPr()
                for old in tcPr.findall(qn("w:tcW")):
                    tcPr.remove(old)
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(int(widths_in[j] * 1440)))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)


def _repeat_header_row(row):
    """让表头行在表格跨页时,每一页顶部都重复显示(w:tblHeader)。

    Word 默认只在第一页显示表头,跨页后续页没有列名、读者对不上列。
    在表头行的 trPr 里加 <w:tblHeader/>,Word 会在每个分页处自动重绘表头。
    """
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:tblHeader")):
        trPr.remove(old)
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def _no_row_split(row):
    """禁止单行内容跨页断开(w:cantSplit),避免一行字被劈成上下两页。"""
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:cantSplit")):
        trPr.remove(old)
    trPr.append(OxmlElement("w:cantSplit"))


def render_table(doc, rows):
    header = split_row(rows[0])
    data = [split_row(r) for r in rows[2:]]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"  # 基底,边框被下面三线表覆盖
    style_table_threeline(table)
    _set_col_widths(table, _compute_col_widths(header, data, _content_width_in(doc)))
    _repeat_header_row(table.rows[0])  # 表头跨页每页重复显示
    _no_row_split(table.rows[0])
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        _shade_cell(cell, THEME["header_fill"])
        para = cell.paragraphs[0]
        _tight(para, line=1.15)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 表头与正文统一左对齐
        run = para.add_run(h)
        run.bold = True
        set_run_fonts(run, size=Pt(9))
    for drow in data:
        row = table.add_row()
        _no_row_split(row)
        cells = row.cells
        for j, val in enumerate(drow[:len(header)]):
            cells[j].text = ""
            para = cells[j].paragraphs[0]
            _tight(para, line=1.15)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(para, val, size=Pt(9))
    # 表后留间隔(docx 表格自身无段后距)
    sp = doc.add_paragraph()
    _tight(sp)
    sp.paragraph_format.space_after = Pt(6)


def render_echarts_block(doc, idx, code, charts_dir):
    """charts_dir/chart_{idx}.png 存在则嵌图 + 自动编号 caption,否则占位框。"""
    png = os.path.join(charts_dir, f"chart_{idx:02d}.png") if charts_dir else None
    if png and os.path.exists(png):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(png, width=Inches(6.0))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(6)
        r = cap.add_run(f"图 {idx}")
        set_run_fonts(r, zh=ZH_HEADING_FONT, en=ZH_HEADING_FONT, size=Pt(9))
        r.font.color.rgb = CAPTION_GRAY
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"［图表占位 chart_{idx:02d} · ECharts 待 render_echarts.js 渲染］")
        run.italic = True
        run.font.color.rgb = PLACEHOLDER_GRAY
        set_run_fonts(run, size=Pt(9))


# ============================================================
# 图片嵌入:Obsidian ![[name]] / markdown ![](path) → 本地文件 → docx
# ============================================================
_IMG_EXTS = (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp")
_SKIP_DIRS = {".git", ".obsidian", "node_modules", "__pycache__", ".trash"}


def _walk_index(root):
    """递归索引 vault 下所有文件名 → 首个路径(兜底查找)。跳过隐藏/大目录。"""
    idx = {}
    for dp, dns, fns in os.walk(root):
        dns[:] = [d for d in dns if d not in _SKIP_DIRS and not d.startswith(".")]
        for fn in fns:
            idx.setdefault(fn, os.path.join(dp, fn))
    return idx


def resolve_image(item, md_dir, vault_root, cache):
    """图片 embed → 本地绝对路径(不存在返回 None)。"""
    tgt = item["target"]
    if item["kind"] == "md":
        if tgt.startswith("file:///"):
            local = unquote(tgt[len("file:///"):]).replace("/", os.sep)
            return local if os.path.exists(local) else None
        if tgt.startswith(("http://", "https://")):
            return None  # 远程图本骨架不下载
        cand = os.path.join(md_dir, unquote(tgt).replace("/", os.sep))
        return cand if os.path.exists(cand) else None
    # Obsidian wiki embed:按文件名在常见位置找,再 walk 兜底
    name = tgt.replace("/", os.sep)
    base = os.path.basename(name)
    for c in (
        os.path.join(md_dir, name),
        os.path.join(md_dir, base),
        os.path.join(md_dir, "attachments", base),
        os.path.join(vault_root, base),
        os.path.join(vault_root, "attachments", base),
    ):
        if os.path.exists(c):
            return c
    if cache["idx"] is None:
        cache["idx"] = _walk_index(vault_root)
    return cache["idx"].get(base)


def _no_table_border(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def _per_row(n):
    """一行排几张:1→1(大图) 2→2 3→3 4→2(2×2) 5+→3"""
    if n <= 1:
        return 1
    if n in (2, 4):
        return 2
    return 3


def add_image_grid(doc, paths, total_width=6.2):
    """多张图并排嵌入(无边框表格,居中)。"""
    n = len(paths)
    ncol = _per_row(n)
    img_w = Inches(min(total_width / ncol, 5.5))
    table = doc.add_table(rows=0, cols=ncol)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _no_table_border(table)
    for k in range(0, n, ncol):
        row = paths[k:k + ncol]
        cells = table.add_row().cells
        for j in range(ncol):
            para = cells[j].paragraphs[0]
            _tight(para)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if j < len(row):
                try:
                    para.add_run().add_picture(row[j], width=img_w)
                except Exception:
                    r = para.add_run(f"［无法嵌入:{os.path.basename(row[j])}］")
                    r.italic = True
                    r.font.color.rgb = PLACEHOLDER_GRAY
                    set_run_fonts(r, size=Pt(9))
    doc.add_paragraph()  # 图组后留间隔


def add_missing_placeholder(doc, item):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = os.path.basename(item["target"]) if item["target"] else "?"
    r = p.add_run(f"［图片缺失:{label}］")
    r.italic = True
    r.font.color.rgb = PLACEHOLDER_GRAY
    set_run_fonts(r, size=Pt(9))


def render_images(doc, items, md_dir, vault_root, cache):
    """一个图片块:并排嵌入已找到的图,逐个标注缺失的图。"""
    resolved = [(resolve_image(it, md_dir, vault_root, cache), it) for it in items]
    found = [p for p, _ in resolved if p]
    if found:
        add_image_grid(doc, found)
    for p, it in resolved:
        if not p:
            add_missing_placeholder(doc, it)


# ============================================================
# 文档级样式
# ============================================================
def _apply_normal_style(doc):
    """Normal 样式:双字体 + 1.4 倍行距 + 段后 6pt(整篇默认)"""
    style = doc.styles["Normal"]
    style.font.size = BODY_SIZE
    pf = style.paragraph_format
    pf.line_spacing = 1.4
    pf.space_after = Pt(6)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), EN_FONT)
    rfonts.set(qn("w:hAnsi"), EN_FONT)
    rfonts.set(qn("w:eastAsia"), ZH_FONT)


def _add_field(paragraph, instr):
    """往段落里插一个 Word 域(如 PAGE)"""
    r_begin = OxmlElement("w:r")
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r_begin.append(fld)
    r_instr = OxmlElement("w:r")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = f" {instr} "
    r_instr.append(it)
    r_end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_end.append(fld_end)
    for r in (r_begin, r_instr, r_end):
        paragraph._p.append(r)


def setup_header_footer(doc, title):
    """页眉:文档标题(右对齐小字 + 底线);页脚:居中页码"""
    section = doc.sections[0]
    # 页眉
    hp = section.header.paragraphs[0]
    _tight(hp)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if title:
        r = hp.add_run(title)
        set_run_fonts(r, zh=ZH_HEADING_FONT, en=ZH_HEADING_FONT, size=Pt(8))
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    _set_para_borders(hp, {"bottom": (4, "D0D0D0")})
    # 页脚:页码域
    fp = section.footer.paragraphs[0]
    _tight(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run()
    set_run_fonts(r, size=Pt(9))
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    _add_field(fp, "PAGE")


def pick_title(fm_dict, blocks, input_path):
    """文档标题(用于页眉 + 正文大标题):frontmatter title → 首个 H1 → 输入 md 文件名。
    用输入文件名而非输出文件名兜底——输入文件名反映文档内容,输出文件名可能是临时的。"""
    if fm_dict.get("title"):
        return fm_dict["title"]
    for b in blocks:
        if b["type"] == "heading" and b["level"] == 1:
            return b["text"]
    return os.path.splitext(os.path.basename(input_path))[0]


# ============================================================
# equity-research 封面 rating box(从 frontmatter 合成)
# ============================================================
COVER_FIELDS = [
    ("评级", "rating"),
    ("现价", "current_price"),
    ("目标价", "target_price"),
    ("空间", "upside"),
    ("市值", "market_cap"),
    ("行业", "sector"),
    ("分析师", "analyst"),
    ("日期", "date"),
]


def build_cover(doc, fm):
    accent = _rgb(THEME["accent"])
    # 眉标
    p = doc.add_paragraph()
    r = p.add_run("首次覆盖 · INITIATING COVERAGE")
    r.bold = True
    set_run_fonts(r, zh=ZH_HEADING_FONT, en=ZH_HEADING_FONT, size=Pt(11))
    r.font.color.rgb = accent
    _set_para_borders(p, {"bottom": (8, THEME["accent"])})

    # 公司标题
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(fm.get("company", "") or fm.get("title", ""))
    r.bold = True
    set_run_fonts(r, zh=ZH_HEADING_FONT, en=ZH_HEADING_FONT, size=Pt(22))

    # ticker 副行
    if fm.get("ticker"):
        p = doc.add_paragraph()
        r = p.add_run(fm["ticker"])
        set_run_fonts(r, size=Pt(11))
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # rating box(两列三线表)
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    style_table_threeline(t)
    for label, key in COVER_FIELDS:
        val = fm.get(key, "")
        if not val:
            continue
        cells = t.add_row().cells
        _shade_cell(cells[0], THEME["header_fill"])
        para_k = cells[0].paragraphs[0]
        _tight(para_k, line=1.15)
        rk = para_k.add_run(label)
        rk.bold = True
        set_run_fonts(rk, size=Pt(10))
        para_v = cells[1].paragraphs[0]
        _tight(para_v, line=1.15)
        add_inline_runs(para_v, val, size=Pt(10))

    doc.add_paragraph()  # 与正文间隔


# ============================================================
# main
# ============================================================
def main():
    global THEME
    ap = argparse.ArgumentParser(description="Obsidian markdown → docx 导出器")
    ap.add_argument("--input", required=True, help="输入 .md 路径")
    ap.add_argument("--output", required=True, help="输出 .docx 路径")
    ap.add_argument("--charts-dir", default=None, help="已渲染 PNG 目录(chart_NN.png)")
    ap.add_argument("--vault-root", default=None, help="vault 根(图片兜底查找;默认=输入 md 所在目录)")
    ap.add_argument("--profile", default="default", choices=["default", "equity-research"])
    args = ap.parse_args()

    THEME = THEMES[args.profile]

    with open(args.input, encoding="utf-8") as f:
        text = f.read()

    fm, body = strip_frontmatter(text)
    fm_dict = parse_frontmatter(fm)
    blocks = parse_blocks(body)

    doc = Document()
    _apply_normal_style(doc)
    setup_header_footer(doc, pick_title(fm_dict, blocks, args.input))

    # profile = equity-research 时,从 frontmatter 合成封面 rating box
    if args.profile == "equity-research":
        build_cover(doc, fm_dict)
    else:
        # default:文档若自带 H1 大标题就不动;没有则在正文顶部补一个
        # (取 frontmatter title → 文件名,对齐 Obsidian 顶部把文件名显示为最大标题的行为)。
        # 否则 docx 开头直接进正文,像"丢了最大标题"。
        has_h1 = any(b["type"] == "heading" and b["level"] == 1 for b in blocks)
        title = pick_title(fm_dict, blocks, args.input)
        if not has_h1 and title:
            render_heading(doc, 1, title)

    md_dir = os.path.dirname(os.path.abspath(args.input))
    vault_root = os.path.abspath(args.vault_root) if args.vault_root else md_dir
    render_blocks(doc, blocks, charts_dir=args.charts_dir, md_dir=md_dir, vault_root=vault_root)

    doc.save(args.output)
    n_img = sum(len(b["items"]) for b in blocks if b["type"] == "images")
    print(f"OK  blocks={len(blocks)}  echarts={sum(1 for b in blocks if b['type']=='echarts')}  images={n_img}  -> {args.output}")


if __name__ == "__main__":
    main()
